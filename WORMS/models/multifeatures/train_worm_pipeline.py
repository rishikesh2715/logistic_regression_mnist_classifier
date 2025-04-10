# train_worm_pipeline.py
import os, sys, time, datetime, tkinter as tk
from tkinter import filedialog, messagebox
import numpy as np, cv2, matplotlib.pyplot as plt
from tqdm import tqdm
from joblib import dump
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import SGDClassifier
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report, log_loss

from skimage.feature import hog
from skimage.measure import regionprops, label
from skimage.morphology import skeletonize

# Optional: for report generation using docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ------------- Prompt for directories (using Tkinter) -------------
root = tk.Tk()
root.withdraw()
dataset_dir = filedialog.askdirectory(title="Select dataset folder (contains subfolders '0' and '1')")
if not dataset_dir:
    messagebox.showerror("Error", "Dataset folder not selected."); sys.exit(1)
models_base = filedialog.askdirectory(title="Select folder to store models and outputs")
if not models_base:
    messagebox.showerror("Error", "Models folder not selected."); sys.exit(1)
os.makedirs(models_base, exist_ok=True)

# ------------- Preprocessing and Mask Creation -------------
def preprocess_and_mask(gray):
    # Apply median blur
    g = cv2.medianBlur(gray, 3)
    # Canny edge detection and blend with blurred image
    edges = cv2.Canny(g, 50, 150)
    blended = cv2.addWeighted(g, 0.8, edges, 0.2, 0)
    # Otsu thresholding
    _, mask = cv2.threshold(blended, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    # Morphological open and close (1 iteration each)
    se1 = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3, 3))
    mask = cv2.morphologyEx(mask, cv2.MORPH_OPEN, se1, iterations=1)
    se2 = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
    mask = cv2.morphologyEx(mask, cv2.MORPH_CLOSE, se2, iterations=1)
    
    # Contour filtering: relaxed thresholds; keep regions with area >= 100 and eccentricity >= 0.7.
    bool_mask = mask > 0
    lbl = label(bool_mask)
    keep = np.zeros_like(mask, dtype=np.uint8)
    for r in regionprops(lbl):
        if r.area >= 100 and r.eccentricity >= 0.7:
            keep[lbl == r.label] = 255
    return keep

# ------------- Feature Extraction -------------
def extract_features(gray):
    mask = preprocess_and_mask(gray)
    if mask.sum() == 0:
        # Fallback: if mask is empty, use whole image (so features are non-zero)
        mask[:] = 255
    # Compute masked HOG: set background to 0
    masked = gray.copy()
    masked[mask==0] = 0
    hog_vec = hog(masked, orientations=12, pixels_per_cell=(8,8),
                  cells_per_block=(2,2), visualize=False, feature_vector=True).astype(np.float32)
    
    # Extract shape features from the mask via regionprops on (mask > 0)
    props = regionprops((mask>0).astype(np.uint8))
    if props:
        rp = props[0]
        shape_feats = np.array([
            rp.area,
            rp.eccentricity,
            rp.solidity,
            rp.major_axis_length/(rp.minor_axis_length+1e-3)
        ], dtype=np.float32)
    else:
        shape_feats = np.zeros(4, dtype=np.float32)
    return np.concatenate([hog_vec, shape_feats])

# ------------- Dataset Loader -------------
def load_dataset(dataset_dir, test_size=0.2, seed=42):
    X, y = [], []
    for label in [0, 1]:
        folder = os.path.join(dataset_dir, str(label))
        if not os.path.isdir(folder):
            print(f"Folder '{folder}' not found, skipping.")
            continue
        files = [f for f in os.listdir(folder) if f.lower().endswith(".png")]
        for f in tqdm(files, desc=f"Processing class {label}", unit="img"):
            img_path = os.path.join(folder, f)
            img = cv2.imread(img_path, cv2.IMREAD_GRAYSCALE)
            if img is None: continue
            feats = extract_features(img)
            X.append(feats)
            y.append(label)
    X = np.array(X)
    y = np.array(y)
    return train_test_split(X, y, test_size=test_size, random_state=seed, stratify=y)

# ------------- Training function using partial_fit -------------
def train_logistic_regression(X_train, y_train, X_test, y_test, max_epochs=100):
    model = SGDClassifier(loss='log_loss', penalty='l2', max_iter=1, warm_start=True,
                           class_weight='balanced', random_state=42)
    train_losses, test_losses, train_accs, test_accs = [], [], [], []
    classes = np.unique(y_train)
    for epoch in range(1, max_epochs+1):
        t0 = time.time()
        model.partial_fit(X_train, y_train, classes=classes)
        # Get losses using predict_proba
        y_train_proba = model.predict_proba(X_train)
        loss_train = log_loss(y_train, y_train_proba)
        train_losses.append(loss_train)
        y_train_pred = model.predict(X_train)
        acc_train = accuracy_score(y_train, y_train_pred)
        train_accs.append(acc_train)
        y_test_proba = model.predict_proba(X_test)
        loss_test = log_loss(y_test, y_test_proba)
        test_losses.append(loss_test)
        y_test_pred = model.predict(X_test)
        acc_test = accuracy_score(y_test, y_test_pred)
        test_accs.append(acc_test)
        epoch_time = time.time() - t0
        print(f"Epoch {epoch:3d}/{max_epochs} - Train Loss: {loss_train:.4f}, Train Acc: {acc_train:.4f} | "
              f"Test Loss: {loss_test:.4f}, Test Acc: {acc_test:.4f} | Time: {epoch_time:.2f}s")
    return model, (train_losses, test_losses, train_accs, test_accs)

# ------------- Main -----------------
def main():
    print("Loading dataset...")
    X_train, X_test, y_train, y_test = load_dataset(dataset_dir, test_size=0.2, seed=42)
    print(f"Dataset loaded: Train samples: {len(X_train)}, Test samples: {len(X_test)}")
    print(f"Feature shape: {X_train.shape} (each sample has {X_train.shape[1]} features)")
    
    # Save a note on input image size (assume all images are same size)
    sample_img = cv2.imread(os.path.join(dataset_dir, "0", os.listdir(os.path.join(dataset_dir, "0"))[0]), cv2.IMREAD_GRAYSCALE)
    img_size = sample_img.shape  if sample_img is not None else "Unknown"
    
    # Scale features
    from sklearn.preprocessing import StandardScaler
    scaler = StandardScaler()
    X_train = scaler.fit_transform(X_train)
    X_test  = scaler.transform(X_test)
    
    print("Training logistic regression using partial_fit...")
    t0 = time.time()
    model, history = train_logistic_regression(X_train, y_train, X_test, y_test, max_epochs=100)
    train_time = time.time() - t0

    y_test_pred = model.predict(X_test)
    acc = accuracy_score(y_test, y_test_pred)
    cm = confusion_matrix(y_test, y_test_pred)
    cr = classification_report(y_test, y_test_pred, digits=4)
    print("Final Test Accuracy:", acc)
    
    # Save training curves
    epochs = range(1, 101)
    train_losses, test_losses, train_accs, test_accs = history
    plt.figure()
    plt.plot(epochs, train_losses, label="Train Loss")
    plt.plot(epochs, test_losses, label="Test Loss")
    plt.xlabel("Epoch")
    plt.ylabel("Loss")
    plt.title("Train vs Test Loss")
    plt.legend()
    loss_path = "training_loss_curve.png"
    plt.savefig(loss_path); plt.close()

    plt.figure()
    plt.plot(epochs, train_accs, label="Train Acc")
    plt.plot(epochs, test_accs, label="Test Acc")
    plt.xlabel("Epoch")
    plt.ylabel("Accuracy")
    plt.title("Train vs Test Accuracy")
    plt.legend()
    acc_path = "training_accuracy_curve.png"
    plt.savefig(acc_path); plt.close()

    # Plot confusion matrix as a heatmap
    plt.figure()
    plt.imshow(cm, interpolation="nearest", cmap=plt.cm.Blues)
    plt.title("Confusion Matrix")
    plt.colorbar()
    tick_marks = np.arange(2)
    plt.xticks(tick_marks, ["0", "1"])
    plt.yticks(tick_marks, ["0", "1"])
    for i in range(2):
        for j in range(2):
            plt.text(j, i, format(cm[i,j], "d"),
                     horizontalalignment="center",
                     color="white" if cm[i,j] > cm.max()/2.0 else "black")
    plt.ylabel("True label"); plt.xlabel("Predicted label")
    plt.tight_layout()
    cm_path = "confusion_matrix.png"
    plt.savefig(cm_path); plt.close()

    # ---------------- Create output folder ---------------- #
    stamp = datetime.datetime.now().strftime("%m%d")
    folder_name = f"worm_{acc*100:.2f}_{stamp}"
    out_folder = os.path.join(models_base, folder_name)
    os.makedirs(out_folder, exist_ok=True)
    dump(model, os.path.join(out_folder, "trained_model.joblib"))
    os.rename(loss_path, os.path.join(out_folder, "training_loss_curve.png"))
    os.rename(acc_path, os.path.join(out_folder, "training_accuracy_curve.png"))
    os.rename(cm_path, os.path.join(out_folder, "confusion_matrix.png"))
    
    # ---------------- Generate documentation ---------------- #
    if DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading("Worm Classification Training Report", 0)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light List'
        
        r = table.add_row().cells
        r[0].text = "Visual Verification of Input Data"
        r[1].text = "[Insert notes on visual inspection of sample images here]"
        
        r = table.add_row().cells
        r[0].text = "Data Splits (Train / Test)"
        r[1].text = f"Train: {len(X_train)}, Test: {len(X_test)} (20% test)"
        
        r = table.add_row().cells
        r[0].text = "Input Image Size"
        r[1].text = str(img_size)
        
        r = table.add_row().cells
        r[0].text = "Image Preprocessing"
        r[1].text = ("Median blur, Canny edge + blend, Otsu threshold, morph open/close (1,1), "
                     "contour filtering (area>=100, ecc>=0.7, fallback to whole image)")
        
        r = table.add_row().cells
        r[0].text = "Logistic Regression Parameters"
        r[1].text = "SGDClassifier(loss='log_loss', penalty='l2', class_weight='balanced')"
        
        r = table.add_row().cells
        r[0].text = "Optimizer and Training"
        r[1].text = "SGD partial_fit over 100 epochs"
        
        r = table.add_row().cells
        r[0].text = "Confusion Matrix (Test)"
        r[1].text = str(cm)
        
        r = table.add_row().cells
        r[0].text = "Execution Times"
        r[1].text = f"Training time: {train_time:.2f} sec"
        
        doc.save(os.path.join(out_folder, "training_report.docx"))
        print("Training report saved to:", os.path.join(out_folder, "training_report.docx"))
    else:
        rep_file = os.path.join(out_folder, "training_report.txt")
        with open(rep_file, "w", encoding="utf-8") as f:
            f.write("Worm Classification Training Report\n")
            f.write("====================================\n\n")
            f.write("A) Training Information\n")
            f.write("i. Visual Verification of Input Data: [placeholder]\n")
            f.write(f"ii. Data Splits: Train {len(X_train)}, Test {len(X_test)} (20% test)\n")
            f.write(f"iii. Input Image Size: {img_size}\n")
            f.write("iv. Image Preprocessing: Median blur, Canny blend, Otsu threshold, morph (1,1), contour filter, fallback\n")
            f.write("v. Logistic Regression Parameters: SGDClassifier(loss='log_loss', penalty='l2', class_weight='balanced')\n")
            f.write("vi. Optimizer: SGD partial_fit, 100 epochs\n\n")
            f.write("B) Confusion Matrix:\n" + str(cm) + "\n\n")
            f.write("C) Training Time:\n" + f"{train_time:.2f} seconds" + "\n\n")
            f.write("Classification Report:\n" + cr + "\n")
        print("Training report saved to:", rep_file)
    
    print("All outputs saved in folder:", out_folder)

if __name__ == "__main__":
    main()
