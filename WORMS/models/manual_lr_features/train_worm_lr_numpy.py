"""
Pure‑NumPy logistic‑regression training for the C‑elegans worm / no‑worm task.

Outputs in the chosen <output‑folder>:
    ├─ worm_lr.npz              – weight vector W and bias b
    ├─ training_report.docx     – (or training_report.txt)
    ├─ training_loss_curve.png
    ├─ training_accuracy_curve.png
    └─ confusion_matrix.png
"""

import os, time, json, datetime, cv2, numpy as np, matplotlib.pyplot as plt
from tqdm import tqdm
from sklearn.metrics import accuracy_score, confusion_matrix, classification_report
from sklearn.model_selection import train_test_split
from skimage.feature import hog, local_binary_pattern
import tkinter as tk
from tkinter.filedialog import askdirectory
from tkinter import messagebox

# optional Word‑doc support
try:
    from docx import Document
    DOCX = True
except ImportError:
    DOCX = False

# ──────────────────────────────────────────────────────────────
# 1.  Feature engineering
# ──────────────────────────────────────────────────────────────
def preprocess(img):
    clahe = cv2.createCLAHE(2.0, (8, 8))
    img = clahe.apply(img)
    img = cv2.medianBlur(img, 3)
    edges = cv2.Canny(img, 50, 150)
    return cv2.addWeighted(img, 0.8, edges, 0.2, 0)

def extract_features(img):
    hog_feat = hog(img, 12, (8, 8), (2, 2), visualize=False, feature_vector=True)
    lbp = local_binary_pattern(img, 8, 1, "uniform")
    hist, _ = np.histogram(lbp, bins=10, range=(0, 10))
    hist = hist.astype(float); hist /= hist.sum() + 1e-7
    return np.concatenate([hog_feat, hist]).astype(np.float32)

# ──────────────────────────────────────────────────────────────
# 2.  NumPy logistic regression
# ──────────────────────────────────────────────────────────────
class LogisticRegressionNP:
    def __init__(self, n_features, lr=0.1):
        rng = np.random.default_rng(42)
        self.W = rng.normal(0, 0.01, size=(n_features,))
        self.b = 0.0
        self.lr = lr

    @staticmethod
    def _sigmoid(z): return 1 / (1 + np.exp(-z))

    def predict_proba(self, X):
        return self._sigmoid(X @ self.W + self.b)

    def predict(self, X):
        return (self.predict_proba(X) >= 0.5).astype(np.int8)

    def _loss(self, y, p):
        eps = 1e-7
        return -(y * np.log(p + eps) + (1 - y) * np.log(1 - p + eps)).mean()

    def fit(self, X, y, X_val, y_val, epochs=60, batch=256):
        N = len(X)
        idx = np.arange(N)
        tr_loss, te_loss, tr_acc, te_acc = [], [], [], []
        for ep in range(1, epochs + 1):
            np.random.shuffle(idx)
            for i in range(0, N, batch):
                j = idx[i:i + batch]
                p = self.predict_proba(X[j])
                grad_W = (X[j].T @ (p - y[j])) / len(j)
                grad_b = (p - y[j]).mean()
                self.W -= self.lr * grad_W
                self.b -= self.lr * grad_b
            # bookkeeping
            p_tr = self.predict_proba(X);   p_te = self.predict_proba(X_val)
            tr_loss.append(self._loss(y, p_tr));      te_loss.append(self._loss(y_val, p_te))
            tr_acc.append(accuracy_score(y, self.predict(X)))
            te_acc.append(accuracy_score(y_val, self.predict(X_val)))
            print(f"Epoch {ep:3d}/60  train-acc {tr_acc[-1]:.3f}  val-acc {te_acc[-1]:.3f}")
        return tr_loss, te_loss, tr_acc, te_acc

# ──────────────────────────────────────────────────────────────
# 3.  I/O helpers
# ──────────────────────────────────────────────────────────────
def prompt_dirs():
    root = tk.Tk(); root.withdraw()
    messagebox.showinfo("Dataset", "Choose dataset folder containing 0 / 1 sub‑folders")
    ds = askdirectory()
    messagebox.showinfo("Save to", "Choose folder to save outputs")
    out = askdirectory(); root.destroy()
    if not ds or not out:
        raise SystemExit("Cancelled.")
    return ds, out

def load_dataset(ds_folder):
    X, y = [], []
    for label in (0, 1):
        folder = os.path.join(ds_folder, str(label))
        files = [f for f in os.listdir(folder) if f.endswith(".png")]
        for f in tqdm(files, desc=f"class {label}", unit="img"):
            g = cv2.imread(os.path.join(folder, f), cv2.IMREAD_GRAYSCALE)
            if g is None: continue
            X.append(extract_features(preprocess(g))); y.append(label)
    return np.array(X), np.array(y, dtype=np.int8), g.shape  # also return img size

# ──────────────────────────────────────────────────────────────
# 4.  Main
# ──────────────────────────────────────────────────────────────
def main():
    ds, out = prompt_dirs()
    X, y, img_sz = load_dataset(ds)
    Xtr, Xte, ytr, yte = train_test_split(
        X, y, test_size=0.2, stratify=y, random_state=42)

    model = LogisticRegressionNP(X.shape[1], lr=0.1)

    t0 = time.time()
    hist = model.fit(Xtr, ytr, Xte, yte, epochs=100)
    train_time = time.time() - t0

    # testing time
    t1 = time.time()
    y_pred = model.predict(Xte)
    test_time = time.time() - t1

    acc = accuracy_score(yte, y_pred)
    cm = confusion_matrix(yte, y_pred)
    class_report = classification_report(yte, y_pred, digits=4)

    # ─── save weights ───────────────────────
    np.savez(os.path.join(out, "worm_lr.npz"),
             W=model.W, b=model.b, feature_dim=X.shape[1])

    # ─── plots ──────────────────────────────
    tr_loss, te_loss, tr_acc, te_acc = hist
    e = np.arange(1, len(tr_loss) + 1)

    plt.figure(); plt.plot(e, tr_loss, label="train"); plt.plot(e, te_loss, label="val")
    plt.title("Binary‑cross‑entropy"); plt.xlabel("epoch"); plt.legend()
    plt.savefig(os.path.join(out, "training_loss_curve.png")); plt.close()

    plt.figure(); plt.plot(e, tr_acc, label="train"); plt.plot(e, te_acc, label="val")
    plt.title("Accuracy"); plt.xlabel("epoch"); plt.legend()
    plt.savefig(os.path.join(out, "training_accuracy_curve.png")); plt.close()

    plt.figure(); plt.imshow(cm, cmap="Blues"); plt.title("Confusion matrix")
    for i in range(2):
        for j in range(2):
            plt.text(j, i, str(cm[i, j]), ha='center', va='center', color='white')
    plt.savefig(os.path.join(out, "confusion_matrix.png")); plt.close()

    # ─── training report ────────────────────
    def add_row(tbl, k, v):
        row = tbl.add_row().cells; row[0].text = k; row[1].text = v

    if DOCX:
        doc = Document(); doc.add_heading("Worm / No‑Worm Training Report", 0)
        tbl = doc.add_table(rows=0, cols=2); tbl.style = "Light List"

        # (A) Training information
        add_row(tbl, "Visual verification of input data",
                "[add your own notes]")
        add_row(tbl, "Data splits",
                f"Train: {len(Xtr)}, Validation: {len(Xte)} (20%)")
        add_row(tbl, "Input image size", f"{img_sz[0]}×{img_sz[1]} pixels")
        add_row(tbl, "Image preprocessing",
                "CLAHE → median‑blur(3) → Canny edges blended (0.8/0.2)")
        add_row(tbl, "Logistic regression parameters",
                f"learning‑rate=0.1, epochs=100, batch=256, "
                f"weights init ~𝒩(0,0.01)")
        add_row(tbl, "Optimizer",
                "Stochastic Gradient Descent (hand‑coded)")

        # (B) Testing information
        add_row(tbl, "Confusion matrix", np.array2string(cm))

        # (C) Execution times
        add_row(tbl, "Training time (s)", f"{train_time:.2f}")
        add_row(tbl, "Testing time  (s)", f"{test_time:.4f}")

        add_row(tbl, "Validation accuracy", f"{acc:.4f}")
        add_row(tbl, "Classification report", class_report)

        doc.save(os.path.join(out, "training_report.docx"))
    else:
        with open(os.path.join(out, "training_report.txt"), "w") as f:
            f.write("Worm / No‑Worm Training Report\n")
            f.write("================================\n\n")
            f.write("A) Training information\n")
            f.write("Visual verification: [add notes]\n")
            f.write(f"Data splits: Train {len(Xtr)}, Validation {len(Xte)}\n")
            f.write(f"Input image size: {img_sz[0]}×{img_sz[1]}\n")
            f.write("Preprocessing: CLAHE, median‑blur, Canny blend\n")
            f.write("Parameters: lr=0.1, epochs=100, batch=256\n")
            f.write("Optimizer: SGD (NumPy implementation)\n\n")
            f.write("B) Confusion matrix\n")
            f.write(f"{cm}\n\n")
            f.write("C) Execution times\n")
            f.write(f"Training: {train_time:.2f}s\nTesting: {test_time:.4f}s\n\n")
            f.write(f"Validation accuracy: {acc:.4f}\n\n")
            f.write(class_report)

    print("\nEverything saved in:", out)

# ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    main()
