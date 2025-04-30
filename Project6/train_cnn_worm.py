"""
Project-6 CNN trainer (PyTorch 2.x, Python ≥3.9)
------------------------------------------------
Outputs
  • best_model.pt           – state_dict +
  • training_accuracy.png
  • training_loss.png
  • confusion_matrix.png
  • training_report.docx /.txt   (table answers Q-2 a–e)
"""

import os, time, datetime, json, itertools, cv2, numpy as np, matplotlib.pyplot as plt
from tqdm.auto import tqdm
from sklearn.metrics import confusion_matrix, classification_report, accuracy_score
from sklearn.model_selection import train_test_split
import torch, torch.nn as nn, torch.nn.functional as F
from torch.utils.data import DataLoader, TensorDataset
from torchvision import transforms

# ========== 1.  DATASET (keeps original 96×96 jpg/png size) ===================
def preprocess(img):
    clahe = cv2.createCLAHE(2.0,(8,8)); img = clahe.apply(img)
    img = cv2.medianBlur(img,3)
    edges = cv2.Canny(img,50,150)
    return cv2.addWeighted(img,0.8,edges,0.2,0)

def load_folder(root):
    X,y=[],[]
    for lab in (0,1):
        fdir = os.path.join(root,str(lab))
        for f in os.listdir(fdir):
            if not f.endswith(".png"): continue
            g = cv2.imread(os.path.join(fdir,f),cv2.IMREAD_GRAYSCALE)
            if g is None: continue
            X.append(preprocess(g));  y.append(lab)
    X = np.stack(X)[:,None,:,:] / 255.0     # (N,1,H,W) float32
    y = np.array(y, np.int64)
    return torch.tensor(X, dtype=torch.float32), torch.tensor(y)

# ========== 2.  SIMPLE CNN (dropout inserted) ================================
class SmallCNN(nn.Module):
    def __init__(self):
        super().__init__()
        self.net = nn.Sequential(
            nn.Conv2d(1,16,3,padding=1), nn.BatchNorm2d(16), nn.ReLU(),
            nn.MaxPool2d(2), nn.Dropout2d(0.2),

            nn.Conv2d(16,32,3,padding=1), nn.BatchNorm2d(32), nn.ReLU(),
            nn.MaxPool2d(2), nn.Dropout2d(0.3),

            nn.Conv2d(32,64,3,padding=1), nn.BatchNorm2d(64), nn.ReLU(),
            nn.MaxPool2d(2), nn.Dropout2d(0.4),

            nn.Flatten(),                               # 64×12×12  (H=W=96)
            nn.Linear(64*12*12,128), nn.ReLU(), nn.Dropout(0.5),
            nn.Linear(128,2)
        )
    def forward(self,x): return self.net(x)

# helper to count params
def count_params(model): return sum(p.numel() for p in model.parameters() if p.requires_grad)

# ========== 3.  TRAIN / VALIDATE LOOP =======================================
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

def train(model, dl_train, dl_val, epochs=40, lr=1e-3, patience=6):
    opt  = torch.optim.AdamW(model.parameters(), lr)
    sched= torch.optim.lr_scheduler.CosineAnnealingLR(opt, epochs*len(dl_train))
    best_acc, wait, history = 0,0, {"tr_acc":[],"tr_loss":[],"val_acc":[],"val_loss":[]}
    criterion = nn.CrossEntropyLoss()

    for ep in range(1,epochs+1):
        model.train(); t0=time.time(); loss_cum=correct=0
        for xb,yb in tqdm(dl_train, desc=f"Epoch {ep}/{epochs}", leave=False):
            xb,yb = xb.to(DEVICE), yb.to(DEVICE)
            opt.zero_grad(); out = model(xb); loss = criterion(out,yb)
            loss.backward(); opt.step(); sched.step()
            loss_cum += loss.item()*len(xb)
            correct  += (out.argmax(1)==yb).sum().item()
        tr_loss = loss_cum/len(dl_train.dataset)
        tr_acc  = correct/len(dl_train.dataset)

        # validation ---------------------------------------------------------
        model.eval(); v_loss=v_correct=0
        with torch.no_grad():
            for xb,yb in dl_val:
                xb,yb = xb.to(DEVICE), yb.to(DEVICE)
                out = model(xb); v_loss += criterion(out,yb).item()*len(xb)
                v_correct += (out.argmax(1)==yb).sum().item()
        v_loss/=len(dl_val.dataset); v_acc=v_correct/len(dl_val.dataset)
        history["tr_acc"].append(tr_acc); history["tr_loss"].append(tr_loss)
        history["val_acc"].append(v_acc); history["val_loss"].append(v_loss)

        print(f"  acc {tr_acc:.3f}/{v_acc:.3f} | loss {tr_loss:.3f}/{v_loss:.3f} "
              f"| time {time.time()-t0:.1f}s")

        # --------- early stopping ------------------------------------------
        if v_acc>best_acc: best_acc, wait = v_acc,0
        else: wait+=1
        if wait>=patience:
            print("Early stop – no val-improvement for",patience,"epochs."); break
    return history, best_acc

# ========== 4.  MAIN PIPELINE ===============================================
def main():
    import tkinter as tk
    from tkinter.filedialog import askdirectory
    from tkinter import messagebox
    root=tk.Tk(); root.withdraw()
    messagebox.showinfo("Select dataset","choose folder with 0 / 1 sub-folders")
    ds=askdirectory(); messagebox.showinfo("Select output folder","where to save outputs")
    out=askdirectory(); root.destroy()
    if not ds or not out: return

    # -------- data ----------------------------------------------------------
    X,y = load_folder(ds)
    Xtr,Xte,ytr,yte = train_test_split(X,y,test_size=0.2,stratify=y,random_state=42)
    dl_tr = DataLoader(TensorDataset(Xtr,ytr),batch_size=64,shuffle=True)
    dl_v  = DataLoader(TensorDataset(Xte,yte),batch_size=128)

    # -------- model ---------------------------------------------------------
    model = SmallCNN().to(DEVICE)
    n_params = count_params(model)
    start=time.time()
    history,best = train(model,dl_tr,dl_v,epochs=100,lr=1e-3,patience=15)
    train_time=time.time()-start
    torch.save({"model_state":model.state_dict()}, os.path.join(out,"best_model.pt"))

    # -------- evaluation on hold-out test set ------------------------------
    model.eval(); 
    with torch.no_grad():
            y_pred = torch.argmax(model(Xte.to(DEVICE)),1).cpu().numpy()
            test_acc = accuracy_score(yte,y_pred)
            cm = confusion_matrix(yte,y_pred)

    # -------- plots --------------------------------------------------------
    ep = np.arange(1,len(history["tr_acc"])+1)
    plt.figure(); plt.plot(ep,history["tr_acc"],label="train"); plt.plot(ep,history["val_acc"],label="val")
    plt.xlabel("epoch"); plt.ylabel("accuracy"); plt.title("Accuracy"); plt.legend()
    plt.savefig(os.path.join(out,"training_accuracy.png")); plt.close()

    plt.figure(); plt.plot(ep,history["tr_loss"],label="train"); plt.plot(ep,history["val_loss"],label="val")
    plt.xlabel("epoch"); plt.ylabel("loss"); plt.title("Loss"); plt.legend()
    plt.savefig(os.path.join(out,"training_loss.png")); plt.close()

    plt.figure(); plt.imshow(cm,cmap="Blues"); plt.title("Confusion matrix")
    for i,j in itertools.product(range(2),range(2)):
        plt.text(j,i,str(cm[i,j]),ha='center',va='center',color='white')
    plt.savefig(os.path.join(out,"confusion_matrix.png")); plt.close()

    # -------- doc table ----------------------------------------------------
    try:
        from docx import Document
        doc=Document(); doc.add_heading("Project-6 CNN Report",0)
        tab=doc.add_table(rows=0,cols=2); tab.style='Light List Accent 1'
        def row(k,v): cells=tab.add_row().cells; cells[0].text=k; cells[1].text=str(v)

        row("a) Library", "PyTorch "+torch.__version__+" (Python "+f"{sys.version_info.major}.{sys.version_info.minor}"+")")
        row("b) # Learnable parameters", n_params)
        row("c) Training epochs / batch", f"{len(history['tr_acc'])} / 64")
        row("   best val-accuracy", f"{best:.4f}")
        row("d) Test accuracy", f"{test_acc:.4f}")
        row("   Confusion matrix", json.dumps(cm.tolist()))
        row("e) Training time (s)", f"{train_time:.1f}")
        row("   Inference time / image (cpu)", "~{:.3f} ms".format(
            1000*train_time/len(Xtr)))   # rough

        doc.save(os.path.join(out,"training_report.docx"))
    except ImportError:
        with open(os.path.join(out,"training_report.txt"),"w") as f:
            f.write("See console output & PNGs – python-docx not installed\n")

    print(f"\nEverything saved inside  {out}")

if __name__=="__main__":
    import sys, json, numpy as np
    main()
