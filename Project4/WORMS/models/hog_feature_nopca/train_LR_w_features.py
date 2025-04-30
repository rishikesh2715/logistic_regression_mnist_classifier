import os
import sys
import time
import datetime
import numpy as np
import cv2
from tqdm import tqdm
from sklearn.linear_model import SGDClassifier
from sklearn.metrics import confusion_matrix, classification_report, accuracy_score, log_loss
from sklearn.model_selection import train_test_split
from joblib import dump
import matplotlib.pyplot as plt
from skimage.feature import hog, local_binary_pattern
import tkinter as tk
from tkinter.filedialog import askdirectory
from tkinter import messagebox

# Optional: for generating a Word doc table that you can save as PDF
# If you don't have python-docx, you can install it or comment out the doc-writing part
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# -------------------------------
# 1) User Prompts for Directories
# -------------------------------
def prompt_directories():
    root = tk.Tk()
    root.withdraw()  # Hide main window

    # Prompt user to select the dataset folder containing subfolders '0' and '1'
    messagebox.showinfo("Select Dataset Folder", "Please select the dataset folder (containing '0' and '1' subfolders).")
    dataset_dir = askdirectory(title="Select Dataset Folder")
    if not dataset_dir:
        messagebox.showerror("Error", "No dataset folder selected!")
        exit(1)
    
    # Prompt user to select the folder where the model will be saved
    messagebox.showinfo("Select Models Folder", "Please select the folder where you want to save your model.")
    models_base = askdirectory(title="Select Models Folder")
    if not models_base:
        messagebox.showerror("Error", "No models folder selected!")
        exit(1)
    
    return dataset_dir, models_base

# -------------------------------
# 2) Feature Extraction
# -------------------------------
def preprocess(img):
    """Enhanced preprocessing pipeline"""
    # CLAHE for adaptive contrast
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    img = clahe.apply(img)
    
    # Denoising
    img = cv2.medianBlur(img, 3)
    
    # Edge emphasis
    edges = cv2.Canny(img, 50, 150)
    return cv2.addWeighted(img, 0.8, edges, 0.2, 0)


def extract_features(img):
    """
    Extracts combined HOG + LBP features from a preprocessed image.
    HOG:
      orientations=12
      pixels_per_cell=(8,8)
      cells_per_block=(2,2)
    LBP:
      P=8, R=1, 'uniform', 10-bin histogram
    """
    # --- HOG ---
    hog_features = hog(
        img,
        orientations=12,
        pixels_per_cell=(8, 8),
        cells_per_block=(2, 2),
        visualize=False,
        transform_sqrt=False,  # can set True if you like
        feature_vector=True
    )

    # --- LBP ---
    lbp = local_binary_pattern(img, P=8, R=1, method="uniform")
    # LBP values are in range [0, P+2) for uniform patterns => 0..10
    # We'll do a 10-bin histogram
    lbp_hist, _ = np.histogram(lbp, bins=10, range=(0, 10))
    # Normalize LBP histogram (optional but often helps)
    lbp_hist = lbp_hist.astype("float")
    lbp_hist /= (lbp_hist.sum() + 1e-7)

    # Concatenate
    return np.concatenate([hog_features, lbp_hist]).astype(np.float32)

def load_dataset(dataset_dir, test_size=0.2, random_state=42):
    """
    Loads images from '0' and '1' subfolders, extracts features,
    and returns X_train, X_test, y_train, y_test.
    """
    X = []
    y = []
    classes = [0, 1]
    
    print("\nLoading dataset and extracting features:")
    for label in classes:
        folder = os.path.join(dataset_dir, str(label))
        if not os.path.isdir(folder):
            print(f"[Warning] Folder '{folder}' not found. Skipping.")
            continue
        
        # List the .png files
        files = [f for f in os.listdir(folder) if f.endswith('.png')]
        
        for f in tqdm(files, desc=f"Processing class {label}", unit="img"):
            img_path = os.path.join(folder, f)
            try:
                # Load in grayscale
                img = cv2.imread(img_path, cv2.IMREAD_GRAYSCALE)
                if img is None:
                    raise ValueError("Could not read image.")
                
                # Preprocess
                img = preprocess(img)
                
                # Extract features
                feats = extract_features(img)
                X.append(feats)
                y.append(label)
            except Exception as e:
                print(f"[Error] Skipping file {img_path}: {e}")
    
    X = np.array(X)
    y = np.array(y)
    
    # train/test split
    X_train, X_test, y_train, y_test = train_test_split(X, y,
                                                        test_size=test_size,
                                                        random_state=random_state,
                                                        stratify=y)
    return X_train, X_test, y_train, y_test

# -------------------------------
# 3) Training a Logistic Regression (via SGD) with partial_fit
# -------------------------------
def train_logistic_regression(X_train, y_train, X_test, y_test, max_epochs=10):
    """
    Demonstration: partial_fit approach to track train/test loss, accuracy over epochs.
    """
    # We do logistic regression with saga solver via SGDClassifier
    # 'log' loss => logistic regression
    # partial_fit => can iterate over multiple epochs
    model = SGDClassifier(loss='log_loss', penalty='l2', max_iter=1, warm_start=True, random_state=42)
    
    # to track metrics over epochs
    train_losses = []
    test_losses = []
    train_accs = []
    test_accs = []
    
    # Classes for partial_fit
    classes = np.unique(y_train)
    
    for epoch in range(1, max_epochs + 1):
        t0 = time.time()
        # partial_fit on entire training set (batch approach)
        model.partial_fit(X_train, y_train, classes=classes)
        
        # Evaluate on training
        y_train_pred = model.predict_proba(X_train)
        train_loss = log_loss(y_train, y_train_pred)
        train_losses.append(train_loss)
        
        y_train_label = np.argmax(y_train_pred, axis=1)
        train_acc = accuracy_score(y_train, y_train_label)
        train_accs.append(train_acc)
        
        # Evaluate on test
        y_test_pred = model.predict_proba(X_test)
        test_loss = log_loss(y_test, y_test_pred)
        test_losses.append(test_loss)
        
        y_test_label = np.argmax(y_test_pred, axis=1)
        test_acc = accuracy_score(y_test, y_test_label)
        test_accs.append(test_acc)
        
        epoch_time = time.time() - t0
        print(f"Epoch {epoch}/{max_epochs} - "
              f"Train Loss: {train_loss:.4f}, Train Acc: {train_acc:.4f} | "
              f"Test Loss: {test_loss:.4f}, Test Acc: {test_acc:.4f} | "
              f"Time: {epoch_time:.2f}s")
    
    return model, (train_losses, test_losses, train_accs, test_accs)

# -------------------------------
# 4) Main Function
# -------------------------------
def main():
    # Prompt user for dataset and models folder
    dataset_dir, models_base = prompt_directories()
    
    # Load and split dataset
    X_train, X_test, y_train, y_test = load_dataset(dataset_dir, test_size=0.2, random_state=42)
    print(f"\nDataset loaded.")
    print(f"Train samples: {len(X_train)}, Test samples: {len(X_test)}")
    print(f"Feature shape: {X_train.shape} (each sample has {X_train.shape[1]} features)")

    # Start training logistic regression
    print("\nTraining Logistic Regression with partial_fit across multiple epochs...")
    train_start = time.time()
    model, metrics_history = train_logistic_regression(X_train, y_train, X_test, y_test, max_epochs=50)
    train_end = time.time()
    train_time = train_end - train_start

    # Final evaluation
    y_test_pred = model.predict(X_test)
    test_acc = accuracy_score(y_test, y_test_pred)
    cm = confusion_matrix(y_test, y_test_pred)
    cr = classification_report(y_test, y_test_pred, digits=4)
    
    print("\nFinal Evaluation on Test Set:")
    print(f"Test Accuracy: {test_acc:.4f}")
    print("Confusion Matrix:\n", cm)
    print("Classification Report:\n", cr)

    # Plot training curves
    train_losses, test_losses, train_accs, test_accs = metrics_history
    plt.figure()
    epochs = range(1, len(train_losses) + 1)
    
    # Loss plot
    plt.plot(epochs, train_losses, label="Train Loss")
    plt.plot(epochs, test_losses, label="Test Loss")
    plt.xlabel("Epoch")
    plt.ylabel("Loss")
    plt.title("Training vs. Testing Loss")
    plt.legend()
    plt.savefig("training_loss_curve.png")
    plt.close()
    
    # Accuracy plot
    plt.figure()
    plt.plot(epochs, train_accs, label="Train Accuracy")
    plt.plot(epochs, test_accs, label="Test Accuracy")
    plt.xlabel("Epoch")
    plt.ylabel("Accuracy")
    plt.title("Training vs. Testing Accuracy")
    plt.legend()
    plt.savefig("training_accuracy_curve.png")
    plt.close()
    
    # ---------------------------
    # Create worm_{ACC}_{MMDD} folder
    # ---------------------------
    date_str = datetime.datetime.now().strftime("%m%d")
    acc_str = f"{test_acc*100:.2f}"
    model_folder_name = f"worm_{acc_str}_{date_str}"
    model_folder_path = os.path.join(models_base, model_folder_name)
    os.makedirs(model_folder_path, exist_ok=True)
    
    # Move plots to that folder
    os.rename("training_loss_curve.png", os.path.join(model_folder_path, "training_loss_curve.png"))
    os.rename("training_accuracy_curve.png", os.path.join(model_folder_path, "training_accuracy_curve.png"))
    
    # Save the model
    model_filename = os.path.join(model_folder_path, "trained_model.joblib")
    dump(model, model_filename)
    print(f"\nModel saved to: {model_filename}")
    
    # ---------------------------
    # Generate a doc containing required info
    # ---------------------------
    if DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading("Worm Classification Training Report", 0)
        
        # Key Info Table
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Light List'
        
        # (A) Training Information
        # i. Visual verification of input data (placeholder)
        row_cells = table.add_row().cells
        row_cells[0].text = "Visual Verification of Input Data"
        row_cells[1].text = "[Place any notes on visually checking the images here]"
        
        # ii. Data splits
        row_cells = table.add_row().cells
        row_cells[0].text = "Training/Validation/Testing Splits"
        row_cells[1].text = f"Train size: {len(X_train)}, Test size: {len(X_test)} (20% test)"
        
        # iii. Input image size
        #    Not strictly used here, but you can fill in if you resized images
        row_cells = table.add_row().cells
        row_cells[0].text = "Input Image Size"
        row_cells[1].text = "Same as original or after any resizing (none in this code)."
        
        # iv. Image Preprocessing
        row_cells = table.add_row().cells
        row_cells[0].text = "Image Preprocessing"
        row_cells[1].text = (
            "CLAHE (clipLimit=2.0, tileGridSize=(8,8))\n"
            "Median blur (3x3)\n"
            "Canny edges blended with original (0.8 / 0.2)\n"
        )
        
        # v. Parameters of logistic regression
        row_cells = table.add_row().cells
        row_cells[0].text = "Logistic Regression Parameters"
        row_cells[1].text = "SGDClassifier(loss='log', penalty='l2', max_iter=1, warm_start=True, random_state=42)"
        
        # vi. Optimizer type and corresponding parameters
        row_cells = table.add_row().cells
        row_cells[0].text = "Optimizer Type & Params"
        row_cells[1].text = (
            "Stochastic Gradient Descent with partial_fit\n"
            "Learning rate = default, no explicit schedule\n"
            "max_epochs = 50"
        )
        
        # (B) Testing info as confusion matrix
        row_cells = table.add_row().cells
        row_cells[0].text = "Confusion Matrix"
        row_cells[1].text = str(cm)
        
        # (C) Training and testing times
        row_cells = table.add_row().cells
        row_cells[0].text = "Execution Times"
        row_cells[1].text = f"Training time: {train_time:.2f} sec"
        
        # Additional notes or placeholders
        row_cells = table.add_row().cells
        row_cells[0].text = "Classification Report"
        row_cells[1].text = cr
        
        row_cells = table.add_row().cells
        row_cells[0].text = "Test Accuracy"
        row_cells[1].text = f"{test_acc:.4f}"
        
        # You could add placeholders for "features summary" or any other notes
        row_cells = table.add_row().cells
        row_cells[0].text = "Feature Summary (Optional)"
        row_cells[1].text = (
            f"Feature dimension: {X_train.shape[1]}\n"
            f"Mean of first feature in training set: {np.mean(X_train[:, 0]):.4f}\n"
            "..."
        )
        
        doc_filename = os.path.join(model_folder_path, "training_report.docx")
        doc.save(doc_filename)
        print(f"Training report saved to: {doc_filename}")
    else:
        # If python-docx isn't available, just output a .txt summary
        txt_filename = os.path.join(model_folder_path, "training_report.txt")
        with open(txt_filename, "w", encoding="utf-8") as f:
            f.write("Worm Classification Training Report\n")
            f.write("====================================\n\n")
            f.write("A) Training Information\n")
            f.write("(i) Visual Verification of Input Data: [placeholder]\n")
            f.write(f"(ii) Data Splits: Train size={len(X_train)}, Test size={len(X_test)}\n")
            f.write("(iii) Input Image Size: [Not resized in this code]\n")
            f.write("(iv) Image Preprocessing: CLAHE, median blur, Canny edges\n")
            f.write("(v) Logistic Regression Parameters: SGDClassifier(loss='log', ...)\n")
            f.write("(vi) Optimizer: SGD, partial_fit\n\n")
            f.write(f"B) Confusion Matrix:\n{cm}\n\n")
            f.write(f"C) Training Time: {train_time:.2f} seconds\n\n")
            f.write("Classification Report:\n")
            f.write(f"{cr}\n")
            f.write(f"Test Accuracy: {test_acc:.4f}\n\n")
            f.write("Feature Summary: [placeholder]\n")
        print(f"Training report saved to: {txt_filename}")

    print("\nDone! All outputs have been saved.")

if __name__ == "__main__":
    main()
